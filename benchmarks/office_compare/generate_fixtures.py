from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook


ROOT = Path(__file__).resolve().parent
FIXTURES = ROOT / "fixtures"


def _write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _write_simple_pdf(path: Path, lines: list[str]) -> None:
    stream_lines = ["BT", "/F1 12 Tf", "72 740 Td"]
    for index, line in enumerate(lines):
        if index:
            stream_lines.append("0 -20 Td")
        stream_lines.append(f"({_escape_pdf_text(line)}) Tj")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{number} 0 obj\n".encode("ascii"))
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(payload)


def generate() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    _write_docx(
        FIXTURES / "project_brief.docx",
        "Orion Project Brief",
        [
            "Project name: Orion.",
            "Objective: improve warehouse inventory accuracy through a monthly reconciliation workflow.",
            "Owner: Chen Rui.",
            "Deadline: 2026-09-30.",
            "Approved budget: CNY 580000.",
            "The brief does not contain profit forecasts for 2027.",
        ],
    )
    _write_docx(
        FIXTURES / "inventory_policy.docx",
        "Inventory Replenishment Policy",
        [
            "Routine replenishment is reviewed every Monday.",
            "Emergency replenishment approver: Liu Fang.",
            "An emergency request must include the SKU, requested quantity, and reason.",
        ],
    )
    _write_simple_pdf(
        FIXTURES / "travel_policy.pdf",
        [
            "Business Travel Policy",
            "Hotel reimbursement limit per night: CNY 680.",
            "Train tickets must use second class seats unless approved in advance.",
        ],
    )
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "销售数据"
    sheet.append(["区域", "销售额", "增长率"])
    sheet.append(["华东", 1800, 0.18])
    sheet.append(["华南", 800, -0.05])
    sheet.append(["东北", 950, 0.03])
    workbook.save(FIXTURES / "sales.xlsx")


if __name__ == "__main__":
    generate()
