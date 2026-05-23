from pathlib import Path
from dataclasses import dataclass
import re

import pandas as pd
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass(frozen=True)
class PdfPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class PdfTable:
    label: str
    caption: str
    page_number: int
    rows: list[list[str]]
    raw_text: str
    extraction_method: str
    confidence: float | None = None


# 根据文件后缀选择合适的文本抽取方式。
# txt/md 直接读文本，csv 转成 markdown 表格，Excel/Word/PDF 调用专门函数；
# 如果遇到未知后缀，就按普通文本尝试读取。
def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".csv":
        return pd.read_csv(path).to_markdown(index=False)
    if suffix in {".xlsx", ".xlsm"}:
        return _read_excel(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    return path.read_text(encoding="utf-8", errors="ignore")


# 读取 Word docx 文件内容。
# 不只读取普通段落，也会遍历表格，把每行单元格用竖线拼起来，
# 最后合并成一段纯文本供检索和 Agent 阅读。
def _read_docx(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# 读取 Excel 工作簿内容。
# 遍历每个工作表和每一行，把非空单元格转成字符串；
# 每个工作表会带上工作表标题，方便后续知道内容来自哪个 sheet。
def _read_excel(path: Path) -> str:
    workbook = load_workbook(path, data_only=True, read_only=True)
    sheets: list[str] = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            sheets.append(f"工作表：{sheet.title}\n" + "\n".join(rows))
    workbook.close()
    return "\n\n".join(sheets)


# 读取 PDF 文件内容。
# 使用 pypdf 逐页抽取文本，并用换行连接；
# 如果某一页抽不到文字，就用空字符串占位，避免程序报错。
def _read_pdf(path: Path) -> str:
    pages = extract_pdf_pages(path)
    return "\n\n".join(f"[page={page.page_number}]\n{page.text}" for page in pages)


def extract_pdf_pages(path: Path) -> list[PdfPage]:
    reader = PdfReader(str(path))
    return [
        PdfPage(page_number=index + 1, text=page.extract_text() or "")
        for index, page in enumerate(reader.pages)
    ]


def extract_pdf_tables(path: Path) -> list[PdfTable]:
    tables = _extract_tables_with_pdfplumber(path)
    if tables:
        return tables
    return _extract_tables_from_text_pages(extract_pdf_pages(path))


def _extract_tables_with_pdfplumber(path: Path) -> list[PdfTable]:
    try:
        import pdfplumber  # type: ignore[import-not-found]
    except ImportError:
        return []

    results: list[PdfTable] = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            labels = _caption_candidates(page_text)
            for table_index, table in enumerate(page.extract_tables() or [], start=1):
                rows = _clean_table_rows(table)
                if not rows:
                    continue
                label, caption = _label_for_table(labels, table_index)
                raw_text = "\n".join(" | ".join(row) for row in rows)
                results.append(
                    PdfTable(
                        label=label,
                        caption=caption,
                        page_number=page_index,
                        rows=rows,
                        raw_text=raw_text,
                        extraction_method="pdfplumber",
                        confidence=0.9,
                    )
                )
    return results


def _extract_tables_from_text_pages(pages: list[PdfPage]) -> list[PdfTable]:
    tables: list[PdfTable] = []
    for page in pages:
        lines = [line.strip() for line in page.text.splitlines()]
        for index, line in enumerate(lines):
            match = _TABLE_CAPTION_RE.match(line)
            if not match:
                continue
            body = _collect_table_like_lines(lines[index + 1 :])
            if not body:
                continue
            label = _normalize_table_label(match.group(1), match.group(2))
            caption = line
            rows = [_split_table_line(item) for item in body]
            tables.append(
                PdfTable(
                    label=label,
                    caption=caption,
                    page_number=page.page_number,
                    rows=rows,
                    raw_text="\n".join(body),
                    extraction_method="pypdf_text",
                    confidence=0.45,
                )
            )
    return tables


_TABLE_CAPTION_RE = re.compile(
    r"^\s*((?:table)|(?:表格?)|(?:TABLE))\s*([0-9]+|[ivxlcdmIVXLCDM]+)\s*[:：.\-]?\s*(.*)$",
    re.IGNORECASE,
)


def _caption_candidates(text: str) -> list[tuple[str, str]]:
    candidates: list[tuple[str, str]] = []
    for line in text.splitlines():
        match = _TABLE_CAPTION_RE.match(line.strip())
        if match:
            candidates.append((_normalize_table_label(match.group(1), match.group(2)), line.strip()))
    return candidates


def _label_for_table(captions: list[tuple[str, str]], table_index: int) -> tuple[str, str]:
    if table_index <= len(captions):
        return captions[table_index - 1]
    label = f"Table {table_index}"
    return label, label


def _normalize_table_label(prefix: str, value: str) -> str:
    normalized_prefix = "Table" if prefix.lower().startswith("table") else "表"
    return f"{normalized_prefix} {value.upper() if value.isalpha() else value}"


def _collect_table_like_lines(lines: list[str], max_lines: int = 80) -> list[str]:
    collected: list[str] = []
    blank_count = 0
    for line in lines[:max_lines]:
        if not line:
            blank_count += 1
            if collected and blank_count >= 2:
                break
            continue
        if _TABLE_CAPTION_RE.match(line) and collected:
            break
        if collected and _looks_like_section_heading(line):
            break
        if _looks_like_table_line(line):
            collected.append(line)
            blank_count = 0
        elif collected:
            break
    return collected


def _looks_like_table_line(line: str) -> bool:
    if "|" in line or "\t" in line:
        return True
    parts = re.split(r"\s{2,}", line.strip())
    has_number = bool(re.search(r"\d", line))
    return len(parts) >= 2 and has_number


def _looks_like_section_heading(line: str) -> bool:
    if len(line) > 100:
        return False
    return bool(re.match(r"^([0-9]+\.|[A-Z][A-Z\s]{4,})", line))


def _split_table_line(line: str) -> list[str]:
    if "|" in line:
        return [part.strip() for part in line.split("|")]
    if "\t" in line:
        return [part.strip() for part in line.split("\t")]
    return [part.strip() for part in re.split(r"\s{2,}", line.strip()) if part.strip()]


def _clean_table_rows(rows: list[list[str | None]]) -> list[list[str]]:
    cleaned: list[list[str]] = []
    for row in rows:
        values = ["" if value is None else str(value).strip() for value in row]
        if any(values):
            cleaned.append(values)
    return cleaned
