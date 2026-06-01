from datetime import datetime
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from uuid import uuid4

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.formula.translate import Translator
from openpyxl.styles import PatternFill
from openpyxl.utils.cell import get_column_letter, range_boundaries

from app.core.config import settings


# 根据文本内容生成 Word 文档。
# filename 用来生成安全的文件名，content 按行拆分后逐段写入 Word；
# 如果 content 没有可用行，就把原始内容作为一个段落写进去。
# 最后返回生成文件在本地磁盘上的 Path。
def generate_word(filename: str, content: str) -> Path:
    settings.artifact_dir.mkdir(parents=True, exist_ok=True)
    path = settings.artifact_dir / f"{_safe_stem(filename)}-{uuid4().hex[:8]}.docx"
    document = Document()
    for block in content.splitlines():
        text = block.strip()
        if text:
            document.add_paragraph(text)
    if not document.paragraphs:
        document.add_paragraph(content)
    document.save(path)
    return path


# 根据文本内容生成 Excel 文件。
# filename 用来生成安全的文件名，content 会按行拆分；
# 每一行再按制表符、逗号或竖线拆成单元格，写入 Sheet1。
# 最后返回生成文件在本地磁盘上的 Path。
def generate_excel(
    filename: str,
    content: str,
    highlight_gt: float | None = None,
    highlight_lt: float | None = None,
    highlight_column: str | None = None,
    highlight_scope: str = "row",
) -> Path:
    settings.artifact_dir.mkdir(parents=True, exist_ok=True)
    path = settings.artifact_dir / f"{_safe_stem(filename)}-{uuid4().hex[:8]}.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    highlight_fill = PatternFill(fill_type="solid", fgColor="FFFF9999")
    rows = [line for line in content.splitlines() if line.strip()]
    parsed_rows = [_split_row(line) for line in (rows or [content])]
    highlight_columns = _resolve_highlight_columns(parsed_rows, highlight_column)
    for row_index, cells in enumerate(parsed_rows, start=1):
        numeric_matches: list[int] = []
        for col_index, value in enumerate(cells, start=1):
            sheet.cell(row=row_index, column=col_index, value=value)
            if highlight_columns is not None and col_index not in highlight_columns:
                continue
            numeric_value = _parse_number(value)
            if _matches_highlight_rule(numeric_value, highlight_gt=highlight_gt, highlight_lt=highlight_lt):
                numeric_matches.append(col_index)
        if numeric_matches:
            if highlight_scope == "cell":
                for col_index in numeric_matches:
                    sheet.cell(row=row_index, column=col_index).fill = highlight_fill
            else:
                for col_index in range(1, len(cells) + 1):
                    sheet.cell(row=row_index, column=col_index).fill = highlight_fill
    workbook.save(path)
    return path


def edit_excel_cells(source_path: str | Path, filename: str, updates: list[dict]) -> Path:
    settings.artifact_dir.mkdir(parents=True, exist_ok=True)
    path = settings.artifact_dir / f"{_safe_stem(filename)}-{uuid4().hex[:8]}.xlsx"
    shutil.copy2(source_path, path)
    workbook = load_workbook(path)
    for update in updates:
        sheet_name = str(update.get("sheet_name") or "").strip()
        sheet = workbook[sheet_name] if sheet_name else workbook.active
        sheet[str(update["cell"]).strip()] = _normalize_excel_value(update.get("value"))
    workbook.save(path)
    workbook.close()
    recalculate_excel_file(path)
    return path


def write_excel_range(
    source_path: str | Path,
    filename: str,
    start_cell: str,
    values: list[list[object]],
    sheet_name: str | None = None,
) -> Path:
    min_column, min_row, _, _ = range_boundaries(start_cell)
    updates = []
    for row_offset, row in enumerate(values):
        for column_offset, value in enumerate(row):
            updates.append(
                {
                    "sheet_name": sheet_name,
                    "cell": f"{get_column_letter(min_column + column_offset)}{min_row + row_offset}",
                    "value": value,
                }
            )
    return edit_excel_cells(source_path, filename, updates)


def fill_excel_formula(
    source_path: str | Path,
    filename: str,
    cell_range: str,
    formula: str,
    sheet_name: str | None = None,
) -> Path:
    min_column, min_row, max_column, max_row = range_boundaries(cell_range)
    origin = f"{get_column_letter(min_column)}{min_row}"
    updates = []
    for row in range(min_row, max_row + 1):
        for column in range(min_column, max_column + 1):
            coordinate = f"{get_column_letter(column)}{row}"
            updates.append(
                {
                    "sheet_name": sheet_name,
                    "cell": coordinate,
                    "value": Translator(formula, origin=origin).translate_formula(coordinate),
                }
            )
    return edit_excel_cells(source_path, filename, updates)


def recalculate_excel_file(path: str | Path, timeout_seconds: int = 120) -> bool:
    source = Path(path).resolve()
    executable = _find_libreoffice()
    if executable is None or not source.exists():
        return False
    with tempfile.TemporaryDirectory(prefix="longchain-lo-", dir=source.parent) as temp_dir:
        root = Path(temp_dir)
        input_dir = root / "input"
        output_dir = root / "output"
        profile_dir = root / "profile"
        input_dir.mkdir()
        output_dir.mkdir()
        profile_dir.mkdir()
        staged = input_dir / source.name
        shutil.copy2(source, staged)
        completed = subprocess.run(
            [
                str(executable),
                f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
                "--headless",
                "--convert-to",
                "xlsx",
                "--outdir",
                str(output_dir),
                str(staged),
            ],
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        converted = output_dir / f"{staged.stem}.xlsx"
        if completed.returncode != 0 or not converted.exists():
            return False
        shutil.copy2(converted, source)
        return True


def _find_libreoffice() -> Path | None:
    candidates = [
        shutil.which("soffice.com"),
        shutil.which("soffice"),
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return Path(candidate)
    return None


# 生成安全的文件主名。
# 会去掉目录和扩展名，只保留中英文、数字、下划线和短横线等安全字符；
# 同时限制长度，避免文件名过长或为空。
def _normalize_excel_value(value: object) -> object:
    if not isinstance(value, str):
        return value
    text = value.strip()
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}(?:[ T]\d{2}:\d{2}:\d{2})?", text):
        return value
    try:
        return datetime.fromisoformat(text)
    except ValueError:
        return value


def _safe_stem(filename: str) -> str:
    stem = Path(filename).stem or "artifact"
    cleaned = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", stem).strip("_")
    return cleaned[:64] or "artifact"


# 把一行文本拆成 Excel 的单元格列表。
# 优先识别制表符，其次识别英文逗号，再识别竖线；
# 如果都没有，就把整行作为一个单元格。
def _split_row(line: str) -> list[str]:
    if "\t" in line:
        return [item.strip() for item in line.split("\t")]
    if "," in line:
        return [item.strip() for item in line.split(",")]
    if "|" in line:
        return [item.strip() for item in line.split("|")]
    return [line.strip()]


def _parse_number(value: str) -> float | None:
    text = value.strip().replace(",", "")
    if not text:
        return None
    if text.endswith("%"):
        text = text[:-1]
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _matches_highlight_rule(
    value: float | None,
    highlight_gt: float | None = None,
    highlight_lt: float | None = None,
) -> bool:
    if value is None:
        return False
    if highlight_gt is not None and value > highlight_gt:
        return True
    if highlight_lt is not None and value < highlight_lt:
        return True
    return False


def _resolve_highlight_columns(rows: list[list[str]], highlight_column: str | None) -> set[int] | None:
    if not highlight_column:
        return None
    target = highlight_column.strip()
    if not target:
        return None
    if target.isdigit():
        return {int(target)}
    column_letter = _column_letter_to_index(target)
    if column_letter is not None:
        return {column_letter}
    if not rows:
        return set()
    normalized_target = _normalize_header(target)
    header = rows[0]
    matches = {
        index
        for index, value in enumerate(header, start=1)
        if _normalize_header(value) == normalized_target
    }
    return matches


def _column_letter_to_index(value: str) -> int | None:
    letters = value.strip().upper()
    if not re.fullmatch(r"[A-Z]+", letters):
        return None
    index = 0
    for letter in letters:
        index = index * 26 + (ord(letter) - ord("A") + 1)
    return index


def _normalize_header(value: str) -> str:
    return re.sub(r"\s+", "", value.strip().lower())
